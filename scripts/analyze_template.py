#!/usr/bin/env python
# -*- coding: utf-8 -*-
r"""分析 .docx 模板文件，提取格式信息。

两种输出模式：

1. **分析模式**（默认）：输出人类可读的 JSON 格式摘要
   python analyze_template.py <template.docx>

2. **模板模式**（--output-template）：输出符合 schema 的标准模板 JSON，
   可直接存入模板库用于论文生成
   python analyze_template.py <template.docx> --output-template my-template.json

输出 JSON 符合 references/template-schema.md 定义的规范。
"""

import json
import re
import sys
import os
from pathlib import Path
from collections import defaultdict

# 在技能安装目录查找 python-docx / lxml，回退到系统路径
_script_dir = os.path.dirname(os.path.abspath(__file__))
_skill_dir = os.path.dirname(_script_dir)
_temp_paths = [
    os.path.join(_skill_dir, '..', '.temp'),
    os.path.join(_skill_dir, '.temp'),
]
for _p in _temp_paths:
    if os.path.isdir(_p):
        sys.path.insert(0, _p)

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from lxml import etree
except ImportError as e:
    print(f"错误: 缺少依赖库。请运行: pip install --target=.temp python-docx lxml")
    print(f"  详情: {e}")
    sys.exit(1)

# --- OOXML 常量 ---
NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _w(tag):
    """生成 w: 命名空间的 tag。"""
    return f'{{{W}}}{tag}'


# --- 工具函数 ---

def emu_to_mm(emu):
    if emu is None:
        return None
    return round(emu / 36000, 1)


def alignment_to_str(al):
    if al is None:
        return "left"
    return {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }.get(al, "left")


# --- 分析函数 ---

def analyze_page_setup(doc):
    """提取页面设置。"""
    result = {}
    for i, section in enumerate(doc.sections):
        result[f"section_{i}"] = {
            "page_width_mm": emu_to_mm(section.page_width),
            "page_height_mm": emu_to_mm(section.page_height),
            "top_margin_mm": emu_to_mm(section.top_margin),
            "bottom_margin_mm": emu_to_mm(section.bottom_margin),
            "left_margin_mm": emu_to_mm(section.left_margin),
            "right_margin_mm": emu_to_mm(section.right_margin),
            "gutter_mm": emu_to_mm(section.gutter),
        }
    return result


def _extract_para_font(para):
    """从段落的第一个有格式的 run 提取字体信息。"""
    for run in para.runs:
        if run.font.name:
            return {
                "font_ascii": run.font.name,
                "font_east_asia": run.font.name,
                "font_size_pt": round(run.font.size / 12700, 1) if run.font.size else None,
                "bold": run.font.bold,
                "italic": run.font.italic,
            }
    return {}


def _extract_para_format(para):
    """提取段落格式。"""
    pf = para.paragraph_format
    result = {}
    if pf.alignment is not None:
        result["alignment"] = alignment_to_str(pf.alignment)
    if pf.line_spacing is not None:
        result["line_spacing"] = round(pf.line_spacing, 1)
    if pf.space_before is not None:
        result["space_before_pt"] = round(pf.space_before / 12700, 1)
    if pf.space_after is not None:
        result["space_after_pt"] = round(pf.space_after / 12700, 1)
    if pf.first_line_indent is not None:
        result["first_line_indent_mm"] = emu_to_mm(pf.first_line_indent)
    return result


def detect_cover_region(doc):
    """检测封面区域：返回 (start, end) 段落索引（end 不含换页符段落）。"""
    page_break_idx = None
    for i, para in enumerate(doc.paragraphs):
        # 检查段落中是否有换页符
        para_xml = para._element.xml
        if '<w:br w:type="page"/>' in para_xml or '<w:br w:type="page" />' in para_xml:
            page_break_idx = i
            break
        # 也检查 pageBreakBefore
        pPr = para._element.find(_w('pPr'))
        if pPr is not None and pPr.find(_w('pageBreakBefore')) is not None:
            page_break_idx = i
            break

    has_table = len(doc.tables) > 0
    has_image = False
    for para in doc.paragraphs[:20]:
        for run in para.runs:
            if run._element.findall('.//' + _w('drawing')) or run._element.findall('.//' + _w('pict')):
                has_image = True
                break

    # 如果有封面元素（表格、图片）且段落间有换页符，判定为封面结束于换页符段落
    if has_table or has_image:
        if page_break_idx and page_break_idx > 0:
            return {"start": 0, "end": page_break_idx - 1}
        # 无换页符但有封面元素，封面区域为前 20 个段落中首个非封面内容之前
        return {"start": 0, "end": 15}

    return None


def detect_heading_patterns(doc):
    """检测标题层级模式。返回 [(level, style_name, numbering_regex), ...]"""
    patterns = []
    heading_styles = set()

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        sl = style_name.lower()

        # 检测样式名中的 heading 语义
        if 'heading' in sl or 'heading' in sl or 'paperh' in sl:
            heading_styles.add(style_name)

        # 检测编号模式
        text = para.text.strip()
        if not text:
            continue

        # "1 绪论" / "1  绪论"
        if re.match(r'^\d+[\s　]+', text) and len(text) < 30:
            heading_styles.add(style_name)
            if not any(p[0] == 1 for p in patterns):
                patterns.append((1, style_name, r'\d+\s+'))
        # "1.1 背景" / "1.1  背景"
        elif re.match(r'^\d+\.\d+[\s　]+', text) and len(text) < 40:
            heading_styles.add(style_name)
            if not any(p[0] == 2 for p in patterns):
                patterns.append((2, style_name, r'\d+\.\d+\s+'))
        # "1.1.1 子节"
        elif re.match(r'^\d+\.\d+\.\d+[\s　]+', text):
            heading_styles.add(style_name)
            if not any(p[0] == 3 for p in patterns):
                patterns.append((3, style_name, r'\d+\.\d+\.\d+\s+'))

    if not patterns:
        # 无编号模式，尝试按样式检测
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else "Normal"
            sl = style_name.lower()
            level = None
            if 'heading1' in sl or 'h1' in sl or style_name == 'PaperH1':
                level = 1
            elif 'heading2' in sl or 'h2' in sl or style_name == 'PaperH2':
                level = 2
            elif 'heading3' in sl or 'h3' in sl or style_name == 'PaperH3':
                level = 3
            if level and not any(p[0] == level for p in patterns):
                patterns.append((level, style_name, None))

    return patterns


def analyze_style_definition(doc, style_name):
    """从 docx 样式 XML 中提取样式定义。如果样式不存在则返回从段落采样推断的格式。"""
    # 尝试从文档 XML 中找到该样式名称的实际定义
    for para in doc.paragraphs:
        if para.style and para.style.name == style_name and para.text.strip():
            font_info = _extract_para_font(para)
            fmt = _extract_para_format(para)
            result = {
                "font_ascii": font_info.get("font_ascii", "宋体"),
                "font_east_asia": font_info.get("font_east_asia", font_info.get("font_ascii", "宋体")),
                "font_size_pt": font_info.get("font_size_pt", 12),
                "bold": font_info.get("bold", False),
                "italic": font_info.get("italic", False),
            }
            if "alignment" in fmt:
                result["alignment"] = fmt["alignment"]
            if "line_spacing" in fmt:
                result["line_spacing"] = fmt["line_spacing"]
            if "space_before_pt" in fmt:
                result["space_before_pt"] = fmt["space_before_pt"]
            if "space_after_pt" in fmt:
                result["space_after_pt"] = fmt["space_after_pt"]
            # 估算首行缩进字符数
            if "first_line_indent_mm" in fmt and fmt["first_line_indent_mm"]:
                indent_mm = fmt["first_line_indent_mm"]
                result["first_line_indent_chars"] = max(0, round(indent_mm / (result["font_size_pt"] * 0.35)))
            else:
                result["first_line_indent_chars"] = 0
            return result

    # 找不到匹配段落，返回默认值
    return {
        "font_ascii": "Times New Roman",
        "font_east_asia": "宋体",
        "font_size_pt": 12,
        "bold": False,
        "line_spacing": 1.0,
        "first_line_indent_chars": 0,
    }


def build_template_json(doc, filepath, template_id=None):
    """从 .docx 文档构建标准模板 JSON。"""
    # 基础 ID
    if template_id is None:
        template_id = re.sub(r'[^a-z0-9-]', '-', Path(filepath).stem.lower().replace('_', '-'))
        template_id = template_id.strip('-') or "imported-template"

    # 页面设置
    sec = doc.sections[0]
    page_setup = {
        "paper": {
            "width_mm": emu_to_mm(sec.page_width) or 210,
            "height_mm": emu_to_mm(sec.page_height) or 297,
        },
        "margins": {
            "top_mm": emu_to_mm(sec.top_margin) or 25,
            "bottom_mm": emu_to_mm(sec.bottom_margin) or 20,
            "left_mm": emu_to_mm(sec.left_margin) or 25,
            "right_mm": emu_to_mm(sec.right_margin) or 20,
        },
        "gutter_mm": emu_to_mm(sec.gutter) or 0,
        "default_line_spacing": 1.0,
    }

    # 收集所有段落按样式分组
    style_paragraphs = defaultdict(list)
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        if para.text.strip():
            style_paragraphs[style_name].append(para)

    # 构建样式字典
    styles = {}
    important_styles = set()

    # 从标题检测中收集样式
    heading_patterns = detect_heading_patterns(doc)
    heading_style_names = set()
    for level, sname, _ in heading_patterns:
        heading_style_names.add(sname)
        important_styles.add(sname)

    # 封面样式
    cover_styles = set()
    for para in doc.paragraphs[:20]:
        if para.style and para.style.name.startswith("Cover"):
            cover_styles.add(para.style.name)
            important_styles.add(para.style.name)

    # 正文样式（找出现次数最多的非标题非封面样式）
    body_style_name = None
    body_para_count = 0
    for sname, paras in style_paragraphs.items():
        if sname in heading_style_names or sname in cover_styles:
            continue
        if sname == "Normal":
            continue  # Normal 通常是 fallback
        if len(paras) > body_para_count:
            body_para_count = len(paras)
            body_style_name = sname

    if body_style_name:
        important_styles.add(body_style_name)

    # 收集所有需要分析的重要样式
    all_style_names = set(style_paragraphs.keys()) | important_styles

    # 对每个样式提取定义
    for sname in sorted(all_style_names):
        if sname == "Normal" and sname not in important_styles:
            continue
        style_def = analyze_style_definition(doc, sname)
        # 补充不可用的字体信息
        if style_def.get("font_ascii") is None:
            style_def["font_ascii"] = "Times New Roman"
        if style_def.get("font_east_asia") is None:
            style_def["font_east_asia"] = style_def["font_ascii"]
        styles[sname] = style_def

    # 确保有基本的 BodyText 样式
    if "BodyText" not in styles and body_style_name:
        body_def = styles.get(body_style_name, {})
        styles["BodyText"] = body_def.copy()

    # 页面顺序
    cover_region = detect_cover_region(doc)
    has_cover = bool(cover_region) or bool(cover_styles)

    page_order = []
    if has_cover:
        page_order.append({"section": "cover", "page_break_after": True})
    page_order.append({"section": "toc", "page_break_after": True})
    page_order.append({"section": "title_page", "page_break_after": False})
    page_order.append({"section": "abstract_zh", "page_break_after": False})
    page_order.append({"section": "abstract_en", "page_break_after": False})
    page_order.append({"section": "body", "page_break_after": False})
    page_order.append({"section": "references", "page_break_after": False})

    # 封面
    cover = {"enabled": False}
    if has_cover:
        cover = {
            "enabled": True,
            "source_template": None,
            "paragraphs_boundary": cover_region,
            "placeholders": {},
        }

    # 标题层级
    heading_levels = []
    for level, sname, num_pattern in sorted(heading_patterns):
        entry = {
            "level": level,
            "style": sname,
        }
        if num_pattern:
            entry["numbering_pattern"] = num_pattern
        heading_levels.append(entry)

    if not heading_levels:
        # 默认三层
        default_headings = [
            (1, "Heading1", r'\d+\s+'),
            (2, "Heading2", r'\d+\.\d+\s+'),
            (3, "Heading3", r'\d+\.\d+\.\d+\s+'),
        ]
        for level, sname, num_pattern in default_headings:
            if sname not in styles:
                styles[sname] = {
                    "font_ascii": "Times New Roman",
                    "font_east_asia": "仿宋体" if level == 1 else "黑体" if level == 2 else "宋体",
                    "font_size_pt": 14 if level == 1 else 12,
                    "bold": level <= 2,
                    "alignment": "left",
                    "line_spacing": 1.0,
                    "first_line_indent_chars": 0,
                }
            heading_levels.append({"level": level, "style": sname, "numbering_pattern": num_pattern})

    # TOC
    toc = {
        "title_style": "TOC_Title",
        "levels": [
            {"heading_level": 1, "style": "TOC_Level1", "indent_chars": 0},
            {"heading_level": 2, "style": "TOC_Level2", "indent_chars": 2},
        ],
        "tab_stop_cm": 14.5,
    }
    for sname in ["TOC_Title", "TOC_Level1", "TOC_Level2"]:
        if sname not in styles:
            styles[sname] = {
                "font_ascii": "Times New Roman",
                "font_east_asia": "黑体" if "Title" in sname else "宋体",
                "font_size_pt": 14 if "Title" in sname else 12,
                "bold": "Title" in sname,
                "alignment": "center" if "Title" in sname else "left",
                "line_spacing": 1.0,
                "first_line_indent_chars": 2 if "Level2" in sname else 0,
            }

    # title_page
    title_page = {
        "title_cn_style": "Title",
        "info_styles": ["TitleInfo"],
        "info_fields": [{"key": "student", "label": "学生"}, {"key": "advisor", "label": "指导教师"}],
    }
    for sname in ["Title", "TitleInfo"]:
        if sname not in styles:
            styles[sname] = {
                "font_ascii": "Times New Roman",
                "font_east_asia": "黑体" if sname == "Title" else "仿宋体",
                "font_size_pt": 16 if sname == "Title" else 12,
                "bold": sname == "Title",
                "alignment": "center",
                "line_spacing": 1.0,
                "first_line_indent_chars": 0,
            }

    # 确保 BodyText 存在
    if "BodyText" not in styles:
        styles["BodyText"] = {
            "font_ascii": "Times New Roman",
            "font_east_asia": "宋体",
            "font_size_pt": 12,
            "bold": False,
            "alignment": "justify",
            "line_spacing": 1.0,
            "space_before_pt": 0,
            "space_after_pt": 0,
            "first_line_indent_chars": 2,
        }

    # 参考文献
    references = {
        "title_style": heading_levels[0]["style"] if heading_levels else "Heading1",
        "body_style": "Reference",
        "title_text": "参考文献",
    }
    if "Reference" not in styles:
        styles["Reference"] = {
            "font_ascii": "Times New Roman",
            "font_east_asia": "宋体",
            "font_size_pt": 10.5,
            "bold": False,
            "alignment": "left",
            "line_spacing": 1.0,
            "first_line_indent_chars": 0,
        }

    template = {
        "id": template_id,
        "name": Path(filepath).stem,
        "description": f"从 {Path(filepath).name} 导入的模板",
        "source": "imported",
        "page_setup": page_setup,
        "styles": styles,
        "page_order": page_order,
        "cover": cover,
        "toc": toc,
        "title_page": title_page,
        "heading_levels": heading_levels,
        "references": references,
        "word_count_estimate": {"chars_per_page_single_spacing": 1500},
    }

    return template


# --- 兼容旧接口的分析模式 ---

def analyze_styles(doc):
    """分析文档中使用的样式（兼容旧接口）。"""
    format_patterns = {}
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        text = para.text.strip()
        if not text:
            continue
        fonts = []
        for run in para.runs:
            f = {}
            if run.font.name:
                f["font_name"] = run.font.name
            if run.font.size:
                f["font_size_pt"] = round(run.font.size / 12700, 1)
            if run.font.bold is not None:
                f["bold"] = run.font.bold
            if run.font.italic is not None:
                f["italic"] = run.font.italic
            if f:
                fonts.append(f)
        para_fmt = _extract_para_format(para)
        key = style_name
        if key not in format_patterns:
            format_patterns[key] = {
                "style_name": style_name,
                "fonts": fonts[:3] if fonts else [],
                "paragraph_format": para_fmt,
                "sample_text": text[:80],
                "count": 1,
            }
        else:
            format_patterns[key]["count"] += 1

    return dict(sorted(format_patterns.items(), key=lambda x: x[1]["count"], reverse=True))


def analyze_document_structure(doc):
    """分析文档结构（兼容旧接口）。"""
    structure = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else "Normal"
        is_heading = any(kw in style_name.lower() for kw in ["heading", "title", "toc", "paperh"])
        entry = {
            "index": i,
            "style": style_name,
            "text_preview": text[:120],
            "is_heading": is_heading,
        }
        if "heading 1" in style_name.lower() or "heading1" in style_name.lower() or style_name == "PaperH1":
            entry["level"] = 1
        elif "heading 2" in style_name.lower() or "heading2" in style_name.lower() or style_name == "PaperH2":
            entry["level"] = 2
        elif "heading 3" in style_name.lower() or "heading3" in style_name.lower() or style_name == "PaperH3":
            entry["level"] = 3
        elif "title" in style_name.lower():
            entry["level"] = 0
        structure.append(entry)
    return structure[:100]


# --- 入口 ---

def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="分析 .docx 模板文件，提取格式信息或生成标准模板 JSON"
    )
    parser.add_argument("file", help=".docx 模板文件路径")
    parser.add_argument("--output-template", metavar="PATH",
                        help="输出标准模板 JSON（符合 template-schema.md 规范）")
    parser.add_argument("--template-id", metavar="ID",
                        help="模板 ID（仅与 --output-template 配合使用）")
    parser.add_argument("--output", metavar="PATH",
                        help="输出分析结果 JSON（兼容旧接口）")
    args = parser.parse_args()

    filepath = args.file
    if not os.path.exists(filepath):
        print(f"错误: 文件不存在: {filepath}")
        sys.exit(1)

    doc = Document(filepath)

    if args.output_template:
        template = build_template_json(doc, filepath, args.template_id)
        with open(args.output_template, "w", encoding="utf-8") as f:
            json.dump(template, f, ensure_ascii=False, indent=2)
        print(f"✅ 标准模板已保存到: {args.output_template}")
        print(f"   ID: {template['id']}, 样式数: {len(template['styles'])}, "
              f"标题层级: {len(template['heading_levels'])}")
        return

    # 兼容旧接口：输出分析摘要
    result = {
        "file": str(filepath),
        "page_setup": analyze_page_setup(doc),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "styles": analyze_styles(doc),
        "structure": analyze_document_structure(doc),
        "cover_detected": detect_cover_region(doc),
        "heading_patterns": [
            {"level": lv, "style": sn, "numbering_pattern": np}
            for lv, sn, np in detect_heading_patterns(doc)
        ],
    }

    # 表格信息
    tables_info = []
    for i, table in enumerate(doc.tables):
        t_info = {"index": i, "rows": len(table.rows), "cols": len(table.columns), "sample_cells": []}
        for row_idx, row in enumerate(table.rows[:3]):
            for col_idx, cell in enumerate(row.cells[:5]):
                if cell.text.strip():
                    t_info["sample_cells"].append({
                        "row": row_idx, "col": col_idx,
                        "text": cell.text.strip()[:60],
                    })
        tables_info.append(t_info)
    result["tables"] = tables_info

    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"分析结果已保存到: {args.output}")
    else:
        print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
